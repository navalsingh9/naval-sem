"""
export_docx.py — NAVAL-SEM  APA 7th Edition Word Report Generator
==================================================================
Public entry point:
    generate_docx(result, indirect_result=None) -> BytesIO

APA 7th formatting applied:
  • Font  : Times New Roman 12 pt body / 10 pt table text
  • Margins: 1 inch (all sides)
  • Tables : NO vertical borders; bold header row; thin horizontal
             rules above header, below header, below last data row only
  • Table numbering: italic "Table N" label above table,
                     bold title on the next line
  • No colour fills in table cells
"""

from __future__ import annotations

import io
import math
from collections import defaultdict
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═════════════════════════════════════════════════════════════════════════════
#  Generic access helpers
# ═════════════════════════════════════════════════════════════════════════════

_SENTINEL = object()


def _get(obj: Any, *keys: str, default: Any = None) -> Any:
    """
    Try to retrieve a value from *obj* using the given keys, in order.
    For each key try attribute access first, then dict-key access.
    Returns *default* if none of the keys match.
    """
    for key in keys:
        # Attribute access (Pydantic models, dataclasses, plain objects)
        try:
            v = getattr(obj, key, _SENTINEL)
            if v is not _SENTINEL:
                return v
        except Exception:
            pass
        # Dict-key access (serialised payloads)
        try:
            if isinstance(obj, dict) and key in obj:
                return obj[key]
        except Exception:
            pass
    return default


def _first(*vals: Any) -> Any:
    """Return the first non-None value from *vals*."""
    for v in vals:
        if v is not None:
            return v
    return None


# ═════════════════════════════════════════════════════════════════════════════
#  Number / string formatting
# ═════════════════════════════════════════════════════════════════════════════

def _fmt(v: Any, dp: int = 3, na: str = "—") -> str:
    """Format a numeric value to *dp* decimal places, or return *na*."""
    if v is None:
        return na
    try:
        f = float(v)
        if math.isnan(f) or math.isinf(f):
            return na
        return f"{f:.{dp}f}"
    except (TypeError, ValueError):
        return str(v)


def _fmt_p(v: Any, na: str = "—") -> str:
    """Format a p-value using APA conventions (< .001 otherwise 3 d.p.)."""
    if v is None:
        return na
    try:
        p = float(v)
        if math.isnan(p) or math.isinf(p):
            return na
        if p < 0.001:
            return "< .001"
        return f"{p:.3f}"
    except (TypeError, ValueError):
        return str(v)


def _stars(pv: Any) -> str:
    """Return APA significance stars for a p-value."""
    if pv is None:
        return ""
    try:
        p = float(pv)
        if p < 0.001:
            return "***"
        if p < 0.01:
            return "**"
        if p < 0.05:
            return "*"
        return ""
    except (TypeError, ValueError):
        return ""


# ═════════════════════════════════════════════════════════════════════════════
#  APA table-border helpers
# ═════════════════════════════════════════════════════════════════════════════

def _border_el(tag: str, size: int = 4, color: str = "000000") -> Any:
    """Return a single-line border OxmlElement (size in 1/8 pt; 4 ≈ 0.5 pt)."""
    el = OxmlElement(tag)
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:color"), color)
    return el


def _nil_el(tag: str) -> Any:
    """Return a nil (suppressed) border OxmlElement."""
    el = OxmlElement(tag)
    el.set(qn("w:val"), "nil")
    return el


def _set_cell_borders(
    cell: Any,
    top: bool = False,
    bottom: bool = False,
    size: int = 4,
    color: str = "000000",
) -> None:
    """
    Set APA-style borders on a single table cell.
    Vertical borders (left, right, insideH, insideV) are always suppressed.
    Only top/bottom horizontal lines are optionally enabled.
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove stale tcBorders element
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)

    borders = OxmlElement("w:tcBorders")
    borders.append(_border_el("w:top",    size, color) if top    else _nil_el("w:top"))
    borders.append(_border_el("w:bottom", size, color) if bottom else _nil_el("w:bottom"))
    for tag in ("w:left", "w:right", "w:insideH", "w:insideV"):
        borders.append(_nil_el(tag))
    tcPr.append(borders)


def _suppress_table_borders(table: Any) -> None:
    """Zero out all table-level borders so cell-level rules take effect."""
    tbl_el = table._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    tbl_borders = OxmlElement("w:tblBorders")
    for tag in ("w:top", "w:left", "w:bottom", "w:right",
                "w:insideH", "w:insideV"):
        tbl_borders.append(_nil_el(tag))
    tblPr.append(tbl_borders)


def _apply_apa_borders(table: Any, n_rows: int, n_cols: int) -> None:
    """
    Enforce APA three-line table rule:
      • Thin horizontal rule above the header row
      • Thin horizontal rule below the header row
      • Thin horizontal rule below the last data row
    No vertical borders anywhere.
    """
    _suppress_table_borders(table)

    for ri, row in enumerate(table.rows):
        is_header = ri == 0
        is_last   = ri == (n_rows - 1)
        for cell in row.cells:
            _set_cell_borders(
                cell,
                top=is_header,                  # above header
                bottom=(is_header or is_last),   # below header / below last row
            )


# ═════════════════════════════════════════════════════════════════════════════
#  Cell / paragraph helpers
# ═════════════════════════════════════════════════════════════════════════════

_BODY_FONT       = "Times New Roman"
_BODY_PT         = Pt(12)
_TABLE_PT        = Pt(10)

_ALIGN = {
    "left":   WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right":  WD_ALIGN_PARAGRAPH.RIGHT,
}


def _clear_cell_fill(cell: Any) -> None:
    """Remove any background shading from a table cell (APA: no colour fills)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)


def _cell_write(
    cell: Any,
    text: str,
    bold: bool   = False,
    italic: bool = False,
    align: str   = "left",
) -> None:
    """
    Clear the cell, write *text* in Times New Roman 10 pt with the
    requested weight / alignment. Removes any background fill.
    """
    # Clear existing paragraph content
    for para in cell.paragraphs:
        para.clear()

    p   = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name   = _BODY_FONT
    run.font.size   = _TABLE_PT
    run.bold        = bold
    run.italic      = italic
    p.alignment     = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)

    _clear_cell_fill(cell)


# ═════════════════════════════════════════════════════════════════════════════
#  Document-level typographic helpers
# ═════════════════════════════════════════════════════════════════════════════

def _add_table_label(doc: Document, n: int, title: str) -> None:
    """
    Insert two paragraphs above the table per APA 7th:
      Line 1 — italic "Table N"
      Line 2 — bold  "Table title"
    """
    p_label = doc.add_paragraph()
    r_label = p_label.add_run(f"Table {n}")
    r_label.font.name = _BODY_FONT
    r_label.font.size = _BODY_PT
    r_label.italic    = True
    r_label.bold      = False
    p_label.paragraph_format.space_before = Pt(12)
    p_label.paragraph_format.space_after  = Pt(0)

    p_title = doc.add_paragraph()
    r_title = p_title.add_run(title)
    r_title.font.name = _BODY_FONT
    r_title.font.size = _BODY_PT
    r_title.bold      = True
    r_title.italic    = False
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(3)


def _add_apa_note(doc: Document, text: str) -> None:
    """Insert an italicised APA Note paragraph below a table."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = _BODY_FONT
    run.font.size = _TABLE_PT
    run.italic    = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(12)


# ═════════════════════════════════════════════════════════════════════════════
#  Table 1 — Measurement Model
# ═════════════════════════════════════════════════════════════════════════════

def _build_measurement_model(doc: Document, result: Any, table_n: int) -> None:
    """
    Table 1: Construct | Indicator | Loading | AVE | CR | α

    Loadings come from parameters where op == '=~'.
    AVE, composite reliability, and Cronbach's alpha come from result.fit.
    Construct-level reliability stats appear only in the first indicator row
    for each latent variable.
    """
    params = _get(result, "parameters") or []
    fit    = _get(result, "fit")        or {}

    meas_params = [p for p in params if _get(p, "op") == "=~"]
    if not meas_params:
        return

    # Group by latent variable (lhs)
    lv_groups: dict[str, list] = defaultdict(list)
    for p in meas_params:
        lv_groups[_get(p, "lhs", default="?")].append(p)

    # Per-construct reliability dicts
    ave_map = _get(fit, "ave") or {}
    cr_map  = _get(fit, "composite_reliability") or {}
    al_map  = _get(fit, "cronbach_alpha")         or {}
    for d in (ave_map, cr_map, al_map):
        if not isinstance(d, dict):
            d = {}

    headers    = ["Construct", "Indicator", "Loading", "AVE", "CR", "α"]
    n_data_rows = sum(len(inds) for inds in lv_groups.values())
    n_rows      = 1 + n_data_rows

    _add_table_label(doc, table_n,
                     "Measurement Model: Factor Loadings and Reliability Indices")

    table = doc.add_table(rows=n_rows, cols=len(headers))
    table.style = "Table Grid"

    # ── Header row ────────────────────────────────────────────────────────
    for ci, h in enumerate(headers):
        _cell_write(table.cell(0, ci), h, bold=True, align="center")

    # ── Data rows ─────────────────────────────────────────────────────────
    row_idx = 1
    for lv_name, indicators in lv_groups.items():
        ave_str = _fmt(_get(ave_map, lv_name), 3)
        cr_str  = _fmt(_get(cr_map,  lv_name), 3)
        al_str  = _fmt(_get(al_map,  lv_name), 3)

        for i, ind in enumerate(indicators):
            loading = _fmt(
                _first(
                    _get(ind, "std_estimate"),
                    _get(ind, "estimate"),
                    _get(ind, "std.all"),
                ),
                3,
            )
            ind_name = _get(ind, "rhs", default="?")

            _cell_write(table.cell(row_idx, 0), lv_name if i == 0 else "", align="left")
            _cell_write(table.cell(row_idx, 1), ind_name,                   align="left")
            _cell_write(table.cell(row_idx, 2), loading,                    align="center")
            _cell_write(table.cell(row_idx, 3), ave_str if i == 0 else "",  align="center")
            _cell_write(table.cell(row_idx, 4), cr_str  if i == 0 else "",  align="center")
            _cell_write(table.cell(row_idx, 5), al_str  if i == 0 else "",  align="center")
            row_idx += 1

    _apply_apa_borders(table, n_rows, len(headers))
    _add_apa_note(
        doc,
        "Note. Loadings are standardised. "
        "AVE = average variance extracted; "
        "CR = composite reliability; "
        "\u03b1 = Cronbach\u2019s alpha. "
        "AVE > .50 and CR > .70 indicate adequate convergent validity.",
    )


# ═════════════════════════════════════════════════════════════════════════════
#  Table 2 — Discriminant Validity (HTMT + √AVE diagonal)
# ═════════════════════════════════════════════════════════════════════════════

def _coerce_htmt_matrix(htmt: Any) -> dict[str, dict[str, float]]:
    """
    Normalise heterogeneous HTMT representations to a nested dict
    { construct_A: { construct_B: float, … }, … }.

    Handles:
      • dict-of-dicts  { A: { B: v } }
      • wrapped object { "matrix": [ {lv1, lv2, htmt}, … ] }
      • flat list      [ {lv1, lv2, htmt}, … ]
      • Pydantic/object with .matrix attribute
    """
    mat: dict[str, dict[str, float]] = {}
    if htmt is None:
        return mat

    # dict-of-dicts
    if isinstance(htmt, dict):
        first_val = next(iter(htmt.values()), None) if htmt else None
        if isinstance(first_val, dict):
            return {k: dict(v) for k, v in htmt.items()}
        # wrapped { "matrix": [...] }
        entries = htmt.get("matrix")
        if entries is None:
            return mat
    elif isinstance(htmt, list):
        entries = htmt
    else:
        # Try attribute .matrix (Pydantic)
        entries = _get(htmt, "matrix")
        if entries is None:
            return mat

    for entry in (entries or []):
        a = _get(entry, "lv1", "construct_a", default="") or ""
        b = _get(entry, "lv2", "construct_b", default="") or ""
        v = _get(entry, "htmt", "value")
        if a and b and v is not None:
            mat.setdefault(a, {})[b] = float(v)
            mat.setdefault(b, {})[a] = float(v)
    return mat


def _build_discriminant_validity(doc: Document, result: Any, table_n: int) -> None:
    """
    Table 2: HTMT ratios (upper triangle) + √AVE (diagonal).
    HTMT cells > 0.90 are **bold**.

    Source: result.fit.htmt (HTMT) and result.fit.ave (√AVE diagonal).
    """
    fit     = _get(result, "fit") or {}
    ave_map = _get(fit, "ave")  or {}
    htmt    = _get(fit, "htmt")

    if not isinstance(ave_map, dict):
        ave_map = {}

    htmt_mat = _coerce_htmt_matrix(htmt)

    # Construct list = union of HTMT keys and AVE keys, sorted
    constructs = sorted(set(list(htmt_mat.keys()) + list(ave_map.keys())))
    if not constructs:
        return

    n = len(constructs)
    n_rows = 1 + n
    n_cols = 1 + n

    _add_table_label(doc, table_n,
                     "Discriminant Validity: HTMT Ratios and \u221aAVE (Diagonal)")

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # ── Header row ────────────────────────────────────────────────────────
    _cell_write(table.cell(0, 0), "", align="center")
    for ci, name in enumerate(constructs):
        _cell_write(table.cell(0, ci + 1), name, bold=True, align="center")

    # ── Data rows ─────────────────────────────────────────────────────────
    for ri, row_name in enumerate(constructs):
        _cell_write(table.cell(ri + 1, 0), row_name, align="left")

        for ci, col_name in enumerate(constructs):
            cell = table.cell(ri + 1, ci + 1)
            if row_name == col_name:
                # Diagonal: √AVE
                ave_v = ave_map.get(row_name)
                try:
                    text = _fmt(math.sqrt(float(ave_v)), 3) if ave_v is not None else "—"
                except (TypeError, ValueError):
                    text = "—"
                _cell_write(cell, text, align="center")
            elif ci > ri:
                # Upper triangle: HTMT value (bold if > 0.90)
                v = htmt_mat.get(row_name, {}).get(col_name)
                text  = _fmt(v, 3) if v is not None else "—"
                emph  = False
                if v is not None:
                    try:
                        emph = float(v) > 0.90
                    except (TypeError, ValueError):
                        pass
                _cell_write(cell, text, bold=emph, align="center")
            else:
                # Lower triangle: blank
                _cell_write(cell, "", align="center")

    _apply_apa_borders(table, n_rows, n_cols)
    _add_apa_note(
        doc,
        "Note. Diagonal entries show the square root of the average variance "
        "extracted (\u221aAVE). Upper-triangle entries are heterotrait\u2013"
        "monotrait (HTMT) ratios. "
        "Values in bold exceed the .90 discriminant validity threshold "
        "(Henseler et al., 2015).",
    )


# ═════════════════════════════════════════════════════════════════════════════
#  Table 3 — Structural Model
# ═════════════════════════════════════════════════════════════════════════════

def _build_structural_model(doc: Document, result: Any, table_n: int) -> None:
    """
    Table 3: Path | β | t-stat | p-value | 95% CI | f² | R²

    Source: result.parameters where op == '~'.
    Significance stars: * p < .05  ** p < .01  *** p < .001
    """
    params  = _get(result, "parameters") or []
    struct  = [p for p in params if _get(p, "op") == "~"]
    if not struct:
        return

    headers = ["Path", "\u03b2", "t-stat", "p-value", "95% CI", "f\u00b2", "R\u00b2"]
    n_rows  = 1 + len(struct)

    _add_table_label(doc, table_n, "Structural Model: Path Coefficients")

    table = doc.add_table(rows=n_rows, cols=len(headers))
    table.style = "Table Grid"

    # ── Header row ────────────────────────────────────────────────────────
    for ci, h in enumerate(headers):
        _cell_write(table.cell(0, ci), h, bold=True, align="center")

    # ── Data rows ─────────────────────────────────────────────────────────
    for ri, p in enumerate(struct, start=1):
        lhs   = _get(p, "lhs", default="?")
        rhs   = _get(p, "rhs", default="?")
        label = f"{rhs} \u2192 {lhs}"

        beta  = _fmt(
            _first(_get(p, "std_estimate"), _get(p, "estimate"), _get(p, "std.all")),
            3,
        )
        t_val = _fmt(
            _first(_get(p, "z"), _get(p, "t"), _get(p, "z_value"), _get(p, "stat")),
            3,
        )
        pv    = _first(
            _get(p, "pvalue"), _get(p, "p_value"),
            _get(p, "p.value"), _get(p, "p"),
        )
        p_str = _fmt_p(pv) + _stars(pv)

        lo = _first(_get(p, "ci_lower"), _get(p, "ci.lower"), _get(p, "ci_lower_95"))
        hi = _first(_get(p, "ci_upper"), _get(p, "ci.upper"), _get(p, "ci_upper_95"))
        ci = f"[{_fmt(lo, 3)}, {_fmt(hi, 3)}]" if (lo is not None and hi is not None) else "—"

        f2 = _fmt(
            _first(_get(p, "f2"), _get(p, "f_squared"), _get(p, "f2_effect")),
            3,
        )
        r2 = _fmt(
            _first(_get(p, "r_squared"), _get(p, "r2")),
            3,
        )

        _cell_write(table.cell(ri, 0), label,  align="left")
        _cell_write(table.cell(ri, 1), beta,   align="center")
        _cell_write(table.cell(ri, 2), t_val,  align="center")
        _cell_write(table.cell(ri, 3), p_str,  align="center")
        _cell_write(table.cell(ri, 4), ci,     align="center")
        _cell_write(table.cell(ri, 5), f2,     align="center")
        _cell_write(table.cell(ri, 6), r2,     align="center")

    _apply_apa_borders(table, n_rows, len(headers))
    _add_apa_note(
        doc,
        "Note. \u03b2 = standardised path coefficient. "
        "CI = confidence interval based on 5,000 bootstrap samples. "
        "f\u00b2 = Cohen\u2019s effect size (small: .02; medium: .15; large: .35). "
        "R\u00b2 = coefficient of determination for the endogenous construct. "
        "* p < .05. ** p < .01. *** p < .001.",
    )


# ═════════════════════════════════════════════════════════════════════════════
#  Table 4 — Indirect Effects (optional)
# ═════════════════════════════════════════════════════════════════════════════

def _build_indirect_effects(
    doc: Document, indirect_result: Any, table_n: int
) -> None:
    """
    Table 4: Path (X→M→Y) | Indirect β | Boot SE | 95% CI | VAF | Significant

    Included only when indirect_result is not None.
    Source: indirect_result.effects
    """
    effects = _get(indirect_result, "effects") or []
    if not effects:
        return

    headers = [
        "Path (X \u2192 M \u2192 Y)",
        "Indirect \u03b2",
        "Boot SE",
        "95% CI",
        "VAF",
        "Significant",
    ]
    n_rows = 1 + len(effects)

    _add_table_label(doc, table_n,
                     "Indirect Effects and Variance Accounted For (VAF)")

    table = doc.add_table(rows=n_rows, cols=len(headers))
    table.style = "Table Grid"

    # ── Header row ────────────────────────────────────────────────────────
    for ci, h in enumerate(headers):
        _cell_write(table.cell(0, ci), h, bold=True, align="center")

    # ── Data rows ─────────────────────────────────────────────────────────
    for ri, eff in enumerate(effects, start=1):
        x_var = _get(eff, "x", "x_var", default="?") or "?"
        m_var = _get(eff, "m", "m_var", default="?") or "?"
        y_var = _get(eff, "y", "y_var", default="?") or "?"
        path  = f"{x_var} \u2192 {m_var} \u2192 {y_var}"

        ind_b = _fmt(
            _first(
                _get(eff, "indirect"),
                _get(eff, "indirect_effect"),
                _get(eff, "beta"),
                _get(eff, "estimate"),
            ),
            3,
        )
        boot_se = _fmt(
            _first(_get(eff, "boot_se"), _get(eff, "se"), _get(eff, "std_error")),
            3,
        )
        lo = _first(_get(eff, "ci_lower"), _get(eff, "ci_lower_95"), _get(eff, "ci.lower"))
        hi = _first(_get(eff, "ci_upper"), _get(eff, "ci_upper_95"), _get(eff, "ci.upper"))
        ci  = f"[{_fmt(lo, 3)}, {_fmt(hi, 3)}]" if (lo is not None and hi is not None) else "—"
        vaf = _fmt(
            _first(_get(eff, "vaf"), _get(eff, "variance_accounted_for")),
            3,
        )
        sig     = _get(eff, "significant")
        sig_str = ("Yes" if sig else "No") if sig is not None else "—"

        _cell_write(table.cell(ri, 0), path,    align="left")
        _cell_write(table.cell(ri, 1), ind_b,   align="center")
        _cell_write(table.cell(ri, 2), boot_se, align="center")
        _cell_write(table.cell(ri, 3), ci,      align="center")
        _cell_write(table.cell(ri, 4), vaf,     align="center")
        _cell_write(table.cell(ri, 5), sig_str, align="center")

    _apply_apa_borders(table, n_rows, len(headers))
    _add_apa_note(
        doc,
        "Note. Indirect \u03b2 = standardised indirect effect; "
        "Boot SE = bootstrapped standard error. "
        "CI = 95% bias-corrected bootstrap confidence interval "
        "based on 5,000 samples. "
        "VAF = variance accounted for (indirect / total effect). "
        "Significance is assessed by whether the 95% CI excludes zero.",
    )


# ═════════════════════════════════════════════════════════════════════════════
#  Public entry point
# ═════════════════════════════════════════════════════════════════════════════

def generate_docx(
    result: Any,
    indirect_result: Optional[Any] = None,
) -> io.BytesIO:
    """
    Build an APA 7th-edition Word report and return it as a seeked BytesIO.

    Parameters
    ----------
    result          : ModelResult (Pydantic object or equivalent dict).
    indirect_result : IndirectResult or equivalent dict; pass None to omit
                      Table 4.

    Returns
    -------
    io.BytesIO — seek position 0, ready to stream or write.
    """
    doc = Document()

    # ── 1-inch margins (all sides), US Letter page ────────────────────────
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))

    # ── Report title ──────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("NAVAL-SEM \u2014 APA 7th Edition Results Report")
    r_title.font.name = _BODY_FONT
    r_title.font.size = Pt(14)
    r_title.bold      = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)

    # ── Tables ────────────────────────────────────────────────────────────
    tbl_n = 1

    _build_measurement_model(doc, result, tbl_n);   tbl_n += 1
    _build_discriminant_validity(doc, result, tbl_n); tbl_n += 1
    _build_structural_model(doc, result, tbl_n);    tbl_n += 1

    if indirect_result is not None:
        _build_indirect_effects(doc, indirect_result, tbl_n)

    # ── Serialise to buffer ───────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═════════════════════════════════════════════════════════════════════════════
#  Smoke test  (python app/export_docx.py)
# ═════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":

    # --- Minimal stub objects ---

    class _FitStub:
        ave                  = {"Satisfaction": 0.612, "Loyalty": 0.583}
        composite_reliability= {"Satisfaction": 0.860, "Loyalty": 0.831}
        cronbach_alpha       = {"Satisfaction": 0.782, "Loyalty": 0.744}
        htmt                 = {"Satisfaction": {"Loyalty": 0.724},
                                 "Loyalty": {"Satisfaction": 0.724}}

    class _Param:
        def __init__(self, op, lhs, rhs, std_estimate=None, z=None,
                     pvalue=None, ci_lower=None, ci_upper=None,
                     f2=None, r_squared=None):
            self.op           = op
            self.lhs          = lhs
            self.rhs          = rhs
            self.std_estimate = std_estimate
            self.z            = z
            self.pvalue       = pvalue
            self.ci_lower     = ci_lower
            self.ci_upper     = ci_upper
            self.f2           = f2
            self.r_squared    = r_squared

    class _ResultStub:
        fit = _FitStub()
        parameters = [
            _Param("=~", "Satisfaction", "SAT1", std_estimate=0.782),
            _Param("=~", "Satisfaction", "SAT2", std_estimate=0.818),
            _Param("=~", "Loyalty",      "LOY1", std_estimate=0.753),
            _Param("=~", "Loyalty",      "LOY2", std_estimate=0.801),
            _Param("~",  "Loyalty", "Satisfaction",
                   std_estimate=0.553, z=7.21, pvalue=0.0001,
                   ci_lower=0.382, ci_upper=0.724,
                   f2=0.208, r_squared=0.306),
        ]

    buf    = generate_docx(_ResultStub())
    length = len(buf.getvalue())

    print(f"[SMOKE TEST]  generate_docx() → BytesIO length = {length:,} bytes")
    assert length > 0, "generate_docx() returned an empty buffer!"
    print("[SMOKE TEST]  Non-empty BytesIO confirmed. All assertions passed. ✓")
